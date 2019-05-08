"""Custom Invitations as Word Documents

Generate a word document with custom invitations.
"""

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

# create a word document
doc = Document()

# read the guests from the file
guests = []
with open("guests.txt", "r") as guests_file:
    for line in guests_file:
        guests.append(line)

# create custom styles
style = doc.styles.add_style("Invitation", WD_STYLE_TYPE.PARAGRAPH)
font = style.font
font.name = "URW Chancery L"
font.size = Pt(20)

style = doc.styles.add_style("GuestInfo", WD_STYLE_TYPE.CHARACTER)
font = style.font
font.name = "Norasi"
font.size = Pt(18)

# create a page for every guest
for guest in guests:
    p = doc.add_paragraph(
        "It would be a pleasure to have the company of", style="Invitation")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run().add_break()
    p.add_run(guest, style="GuestInfo")

    p = doc.add_paragraph(
        "at 11010 Memory Lane on the Evening of", style="Invitation")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p.add_run().add_break()
    p.add_run("April 1st", style="GuestInfo")
    p.add_run().add_break()

    doc.add_paragraph(
        "at 7 o'clock", style="Invitation").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

# save the document
doc.save("invitations.docx")
